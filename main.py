############################################################
########### TOWER TROUBLESHOOTING DATA PIPELINE ############
############################################################

#NOTE: Please download any input files which are in .doc format as.odt

##############################################
########### PART I: DOC SPLITTING ############
##############################################

import os
import subprocess
from docx import Document
from docx.shared import Inches
import tempfile
from docx.oxml.xmlchemy import BaseOxmlElement
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.text.paragraph import Paragraph

import vertexai
from vertexai import generative_models
from vertexai.generative_models import GenerativeModel
import magic
from PIL import Image
import fitz
import io
from datetime import date

os.environ["GCLOUD_PROJECT"] = "genai-sandbox-389908"
os.environ["GOOGLE_CLOUD_QUOTA_PROJECT"] = "genai-sandbox-389908"
project_id = "genai-sandbox-389908"

jsonl_metadata = []

vertexai.init(project=project_id, location="us-central1")

def extract_first_page(input_pdf_path, output_pdf_path):
    # Open the input PDF file
    input_pdf = fitz.open(input_pdf_path)
    
    # Create a new PDF document for the output
    output_pdf = fitz.open()
    
    # Add the first three pages to the new PDF
    for page_num in range(min(3, input_pdf.page_count)):
        page = input_pdf.load_page(page_num)
        output_pdf.insert_pdf(input_pdf, from_page=page_num, to_page=page_num)
    
    # Save the new PDF to the specified output path
    output_pdf.save(output_pdf_path)
    output_pdf.close()
    input_pdf.close()

def convert_pdf_to_single_jpeg(input_file_path, output_jpeg_path):
    try:
        doc = fitz.open(input_file_path)  # Open the PDF file
        page_images = []

        for page_num in range(len(doc)):
            page = doc.load_page(page_num)  # Load the current page
            pix = page.get_pixmap(dpi=50)  # Render page to an image
            img_data = pix.tobytes("png")  # Convert the image to PNG bytes
            img = Image.open(io.BytesIO(img_data))  # Open the image from bytes
            page_images.append(img)

        # Calculate total height and maximum width
        total_width = max(img.width for img in page_images)
        total_height = sum(img.height for img in page_images)

        # Create new blank image
        merged_image = Image.new("RGB", (total_width, total_height))

        # Paste each page image
        y_offset = 0
        for img in page_images:
            merged_image.paste(img, (0, y_offset))
            y_offset += img.height

        # Save as JPEG
        merged_image.save(output_jpeg_path, "JPEG", quality=95)
        print("PDF conversion to single JPEG successful!")
    except Exception as e:
        print(f"Error during PDF to JPEG conversion: {e}")

def convert_to_jpeg(input_file_path, mime_type, output_jpeg_path):
    if mime_type == "application/pdf":
        convert_pdf_to_single_jpeg(input_file_path, output_jpeg_path)
    else:
        try:
            # Open the image
            with Image.open(input_file_path) as img:
                # Save as JPEG with maximum quality
                img.save(output_jpeg_path, quality=100)
            print("Image Conversion successful!")
        except Exception as e:
            print("Error:", e)

def extract_tool_name(project, filepath):
    print("Checking file: "+filepath)
    # Initiate the Model
    #model = GenerativeModel(model_name="gemini-1.0-pro-vision")
    model = GenerativeModel(model_name="gemini-1.5-pro-001")

    # Query the Model
    intro = """ Your job is to take troubleshooting documents and to return the name of the tool to which the document refers. 
    The tool name is usually found in the beginning of the document, like so: XXXXX Troubleshooting Procedures, where XXXXX is the name of the tool. 
    Only return the name of the tool, do not return any other text.
"""

    # build prompt with examples
    prompt = [intro]
    # check mime tpye of image
    mime = magic.Magic(mime=True)
    file_mime_type = mime.from_file(filepath)
    question_images = []
    allowed_formats = ["image/jpeg", "image/png"]

    if file_mime_type not in allowed_formats:
        pages = convert_to_jpeg(filepath, file_mime_type, "converted_img.jpg")
        question_images = [generative_models.Part.from_image(generative_models.Image.load_from_file("converted_img.jpg"))]
    else:
        question_images = [generative_models.Part.from_image(generative_models.Image.load_from_file("converted_img.jpg"))]

    # add question to prompt
    for question_image in question_images:
        prompt.append(question_image)
    question = """What is the tool name?

"""
    prompt.append(question)
    try:
        # Try to generate content using the model
        response = model.generate_content(prompt)

    except ServiceUnavailable as e:
        # Handle any other exceptions that might occur
        print(f"An unexpected error occurred: {e} Passing the document to the extractor.")
        # Optionally, log the full traceback for debugging
        traceback.print_exc()
        # Return a default response or handle the error as needed
        return "Some problem occurred."
    except Exception as e:
        # Handle any other exceptions that might occur
        print(f"An unexpected error occurred: {e} Passing the document to the extractor.")
        # Optionally, log the full traceback for debugging
        traceback.print_exc()
        # Return a default response or handle the error as needed
        return "Some problem occurred."

    return response.text.strip()

def append_text_to_docx(file_path, text_to_append):
    # Open the existing document
    doc = Document(file_path)

    # Add a new paragraph with the text to append
    doc.add_paragraph(text_to_append)

    # Save the modified document
    doc.save(file_path)

def read_document(file_path):
    doc = Document(file_path)
    return doc


def is_split_point(paragraph):
    """Split criteria"""
#    if paragraph.style.name.startswith('Heading 1') or paragraph.style.name.startswith('Heading 2') or paragraph.style.name.startswith('Heading 3'):
#        return True
    if paragraph.style.name.startswith('Heading 1') or paragraph.style.name.startswith('Heading 2'):
        return True
    return False


def split(doc, cut_idx):
    """Splitting into parts by copying source document and cutting off
    irrelevant data."""
    tmpdocfile = write_tmp_doc(doc)
    second_part = doc
    second_elems = list(second_part.element.body.iterchildren())
    for i in range(0, cut_idx):
        remove_element(second_elems[i])
    first_part = Document(tmpdocfile)
    first_elems = list(first_part.element.body.iterchildren())
    for i in range(cut_idx, len(first_elems)):
        remove_element(first_elems[i])
    tmpdocfile.close()
    return (first_part, second_part)


def remove_element(elem):
    elem.getparent().remove(elem)


def write_tmp_doc(doc):
    tmp = tempfile.TemporaryFile()
    doc.save(tmp)
    tmp.seek(0)
    return tmp


def iterparts(doc_path, skip_first=True, bias=0):
    """Iterate over sub-documents by splitting source document into parts"""
    doc = Document(doc_path)
    counter = 0
    while doc:
        split_elem_idx = -1
        doc_body = doc.element.body
        cutted = [doc, None]
        for idx, elem in enumerate(doc_body.iterchildren()):
            if isinstance(elem, CT_P):
                p = Paragraph(elem, doc)
                if is_split_point(p):
                    if split_elem_idx == -1 and skip_first:
                        split_elem_idx = idx
                    else:
                        cutted = split(doc, idx + bias)  # idx-1 to keep previous paragraph
                        counter += 1
                        break
        yield (counter, cutted[0])
        doc = cutted[1]


def sanitize_filename(filename):
    return "".join([c if c.isalnum() or c in ".-" else "" for c in filename])


def write_sections_to_docx(sections, output_dir, tool_name):
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    for section, doc in sections.items():
        section_path = os.path.join(output_dir, f"{sanitize_filename(section)}.docx")
        doc.save(section_path)
        append_text_to_docx(section_path, f"This document is for troubleshooting tool: {tool_name}")
        jsonl_metadata.append({'document': section_path.split('/')[-2], 'tool': tool_name})


def load_parse_and_convert_document(file_path, output_dir, tool_name):
    sections = {}
    for counter, doc in iterparts(file_path):
        if counter == 0:
            continue  # skip the first part if it is not needed
        section_heading = doc.paragraphs[0].text.strip()
        sections[section_heading] = doc
    write_sections_to_docx(sections, output_dir, tool_name)

def convert_docx_to_pdf(docx_path, output_dir):
    # Convert the .docx file to .pdf using LibreOffice
    subprocess.run(['soffice', '--headless', '--convert-to', 'pdf', '--outdir', output_dir, docx_path])
    # Convert the .pdf file to .png using LibreOffice
    pdf_path = os.path.splitext(docx_path)[0] + '.pdf'
    return pdf_path

def convert_all_docx_to_pdf(output_dir):
    for filename in os.listdir(output_dir):
        if filename.endswith('.docx'):
            docx_path = os.path.join(output_dir, filename)
            convert_docx_to_pdf(docx_path, output_dir)

input_dir = "input"
output_dir = 'output'
input_as_pdfs_dir = 'input_as_pdfs'

#Convert input .doc files to .docx
for filename in os.listdir(input_dir):
    if filename.endswith('.odt'):
        doc_path = os.path.join(input_dir, filename)
        subprocess.run(["soffice", "--headless", "--convert-to", 'docx', '--outdir', input_dir, doc_path])

for filename in os.listdir(input_dir):
    if filename.endswith('.docx'):
        file_path = os.path.join(input_dir, filename)
        trunc_name = file_path[:-5]
        trunc_name = trunc_name.split('/')[-1]
        print(trunc_name)
        subprocess.run(['soffice', '--headless', '--convert-to', 'pdf', '--outdir', input_as_pdfs_dir, file_path])
        extract_first_page(input_as_pdfs_dir + '/' + trunc_name + '.pdf', input_as_pdfs_dir + '/' + trunc_name + '2' + '.pdf')
        tool_name = extract_tool_name(project_id, input_as_pdfs_dir + '/' + trunc_name + '2' + '.pdf')
        #append_text_to_docx(file_path, f"This document is for troubleshooting tool: {tool_name}")
        load_parse_and_convert_document(file_path, output_dir + '/' + trunc_name, tool_name)
        convert_all_docx_to_pdf(output_dir + '/' + trunc_name)
        #jsonl_metadata.append({'document': trunc_name, 'tool': tool_name})


######################################
### PART II: LLM IMAGE DESCRIPTION ###
######################################

def query(project, filepath):
    print("Checking file: "+filepath)
    # Initiate the Model
    #model = GenerativeModel(model_name="gemini-1.0-pro-vision")
    model = GenerativeModel(model_name="gemini-1.5-pro-001")

    # Query the Model
    intro = """ You are an expert interpreter of instructions. 
    Your job is to take in a document which contains a mix of text and images. 
    The images are diagrams which you must analyze and provide detailed explanations of (in simple language). 
    If the diagrams contain instructions, make sure to provide the instructions. 
    If the images are tables, explain the content of the tables. 
    Do NOT be lazy - you MUST describe the entire diagram.

    At the beginning of each image description, you MUST state which image you are describing.

    If the document contains multiple images, then describe them all, and clearly state which image you are describing.

    If the document does not contain any images, then please response ONLY with the phrase:
    This document does not contain any diagrams.
"""

    # build prompt with examples
    prompt = [intro]
    # check mime tpye of image
    mime = magic.Magic(mime=True)
    file_mime_type = mime.from_file(filepath)
    question_images = []
    allowed_formats = ["image/jpeg", "image/png"]

    if file_mime_type not in allowed_formats:
        pages = convert_to_jpeg(filepath, file_mime_type, "converted_img.jpg")
        question_images = [generative_models.Part.from_image(generative_models.Image.load_from_file("converted_img.jpg"))]
    else:
        question_images = [generative_models.Part.from_image(generative_models.Image.load_from_file("converted_img.jpg"))]

    # add question to prompt
    for question_image in question_images:
        prompt.append(question_image)
    question = """Your answer:

"""
    prompt.append(question)
    try:
        # Try to generate content using the model
        response = model.generate_content(prompt)

    except ServiceUnavailable as e:
        # Handle any other exceptions that might occur
        print(f"An unexpected error occurred: {e} Passing the document to the extractor.")
        # Optionally, log the full traceback for debugging
        traceback.print_exc()
        # Return a default response or handle the error as needed
        return "Some problem occurred."
    except Exception as e:
        # Handle any other exceptions that might occur
        print(f"An unexpected error occurred: {e} Passing the document to the extractor.")
        # Optionally, log the full traceback for debugging
        traceback.print_exc()
        # Return a default response or handle the error as needed
        return "Some problem occurred."

    return response.text.strip()

final_output_dir = output_dir + '/' + 'final_output'
file_num = 0

# Loop through all sub-directories in the root directory
for subdir, _, files in os.walk(output_dir):
    print(f"Directory: {subdir}")
    if subdir == final_output_dir or subdir == output_dir:
        continue
    # Loop through all files in the current sub-directory
    for file in files:
        file_path = os.path.join(subdir, file)
        print(f"  File: {file_path}")
        if file_path.endswith('.pdf'):
            response = query(project_id, file_path)
            docx_path = file_path[:-4] + ".docx"
            print(f"File path: {file_path}")
            print(f"docx path: {docx_path}")
            if docx_path.split('/')[-1] == ".docx.docx":
                continue
            append_text_to_docx(docx_path, response)
            append_text_to_docx(docx_path, f"This information is taken from troubleshooting document: {subdir.split('/')[-1]}")
            old_output_path = os.path.join(final_output_dir, docx_path.split('/')[-2]).replace(" ", "_")
            convert_docx_to_pdf(docx_path, old_output_path)
            source_pdf = os.path.join(old_output_path, docx_path.split('/')[-1][:-4] + 'pdf')
            #target_pdf = os.path.join(old_output_path, docx_path.split('/')[-2] + '_' + docx_path.split('/')[-1][:-4] + 'pdf')
            target_pdf = os.path.join(old_output_path, str(file_num) + '_' + docx_path.split('/')[-1][:-4] + 'pdf')
            target_pdf = target_pdf.replace(' ', '_')
            print(f"Source: {source_pdf}")
            print(f"Target: {target_pdf}")
            print(f"Renaming {source_pdf} to {target_pdf}")
            if os.path.exists(source_pdf):
                os.rename(source_pdf, target_pdf)
            else:
                print(f"Source file does not exist: {source_pdf}")
    file_num += 1
